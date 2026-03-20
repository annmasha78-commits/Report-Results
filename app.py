import streamlit as st
import os
import io
import platform
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import pytesseract
from groq import Groq

# --- 1. SECURITY & PATH FIX ---
load_dotenv()
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY") or os.getenv("GROQ_API_KEY")

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# --- 2. LOGIC: Smart Extraction (PDF, Word, Images) ---
@st.cache_data(show_spinner=False)
def get_text_from_files(file_list):
    all_text = ""
    for file in file_list:
        try:
            # A. Agar PDF file hai
            if file.name.endswith(".pdf"):
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    all_text += page.extract_text() + "\n"
            
            # B. Agar Word file evidence mein di hai
            elif file.name.endswith(".docx"):
                doc = Document(file)
                for para in doc.paragraphs:
                    all_text += para.text + "\n"
            
            # C. Agar Screenshots/Images hain (OCR)
            elif file.type in ["image/png", "image/jpeg", "image/jpg"]:
                img = Image.open(file)
                ocr_text = pytesseract.image_to_string(img)
                all_text += f"\n[Screenshot Content: {file.name}]\n{ocr_text}\n"
                
        except Exception as e:
            st.error(f"Error reading {file.name}: {e}")
    return all_text

def process_with_ai_batched(prompt_text, api_key_to_use):
    temp_client = Groq(api_key=api_key_to_use)
    try:
        # AI ko mazeed behtar hidayat di hain taake report ki tarteeb kharab na ho
        system_msg = (
            "You are an Elite Cyber Security Analyst. Your task is to analyze the provided screenshots, "
            "PDF data, and logs. Extract critical vulnerabilities (OWASP Top 10 focus). "
            "Keep the tone professional and maintain the technical structure of a standard pentest report. "
            "Do not add conversational filler. Provide only the findings."
        )
        
        response = temp_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": f"Analyze this evidence and provide detailed findings: \n{prompt_text}"}
            ],
            temperature=0.1
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI Analysis Failed: {str(e)}"

def fill_template(template_file, findings):
    doc = Document(template_file)
    # Ye step aapke template ki settings aur format ko barqarar rakhta hai
    for p in doc.paragraphs:
        if "{{RESULT}}" in p.text:
            p.text = p.text.replace("{{RESULT}}", findings)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. UI SETUP ---
st.set_page_config(page_title="CyberReport AI Pro", page_icon="🛡️", layout="wide")

st.title("🛡️ Cyber Pentest Report AI (Pro)")
st.markdown("Automate your reporting by analyzing Screenshots, PDFs, and Docs.")

with st.sidebar:
    st.header("⚙️ Configuration")
    user_api_key = st.text_input("Enter Groq API Key", type="password", value=GROQ_API_KEY if GROQ_API_KEY else "")
    st.divider()
    st.markdown("### How to use:")
    st.write("1. Upload your format in **Template**.")
    st.write("2. Upload all findings in **Evidence**.")
    st.write("3. AI will auto-map everything.")

# Columns update for better file handling
col1, col2 = st.columns(2)
with col1:
    st.subheader("📄 Report Template")
    # Yahan sirf Word Template allow hai kyunke report wahan likhni hai
    template_file = st.file_uploader("Upload Word Template (.docx)", type=["docx"], key="tpl")

with col2:
    st.subheader("📸 Evidence & Data")
    # Yahan sab allow hai: PDF, Images aur purani Word Reports
    evidence_files = st.file_uploader("Upload Evidence (PDF, Images, Docx)", 
                                     type=["pdf", "png", "jpg", "jpeg", "docx"], 
                                     accept_multiple_files=True, key="evd")

# --- 4. EXECUTION ---
if st.button("🚀 Analyze Evidence & Generate Report"):
    if template_file and evidence_files and user_api_key:
        with st.spinner("Extracting data from files and screenshots..."):
            # Step A: Sab files se text nikalna
            full_data = get_text_from_files(evidence_files)
            
            if full_data.strip() == "":
                st.error("No text could be extracted from the uploaded files.")
            else:
                # Step B: AI Analysis
                ai_findings = process_with_ai_batched(full_data, user_api_key)
                
                # Step C: Template Filling
                final_report = fill_template(template_file, ai_findings)
                
                st.success("✅ Report Prepared Successfully!")
                
                st.download_button(
                    label="📥 Download Professional Report",
                    data=final_report,
                    file_name="Final_Pentest_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                with st.expander("Review AI Analysis Results"):
                    st.markdown(ai_findings)
    else:
        st.warning("Please upload the Template, Evidence, and provide an API Key.")
